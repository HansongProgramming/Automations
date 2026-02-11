"""Verify that conditional sections were removed from generated documents"""
from docx import Document
import re
from pathlib import Path

# Check the first generated document
doc_path = 'test_claim_letters/RICHARD_DAVIES_CAPITAL_ONE__EUROPE__PLC_LOC.docx'

if not Path(doc_path).exists():
    print(f"Document not found: {doc_path}")
    exit(1)

doc = Document(doc_path)

print("="*70)
print("VERIFYING SECTION REMOVAL")
print("="*70)
print(f"Checking: {doc_path}\n")

# Look for section markers and delete instructions
sections_found = {}
delete_markers_found = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Check for delete markers
    if '{*Delete' in text:
        delete_markers_found.append((i, text[:100]))
    
    # Check for section markers (22.1, 22.2, 23.1-23.6)
    match = re.match(r'^\s*(2[23]\.\d+)\s+(.+)', text)
    if match:
        section_num = match.group(1)
        section_title = match.group(2)[:50]
        sections_found[section_num] = (i, section_title)

# Expected sections to be REMOVED (no data): 22.1, 22.2, 23.1, 23.4, 23.5, 23.6
# Expected sections to be KEPT (has data): 23.2, 23.3

should_be_removed = ['22.1', '22.2', '23.1', '23.4', '23.5', '23.6']
should_be_kept = ['23.2', '23.3']

print("Delete Markers Found:")
if delete_markers_found:
    print("  ❌ ISSUE: Delete markers should be removed!")
    for idx, text in delete_markers_found:
        print(f"    Line {idx}: {text}")
else:
    print("  ✓ No delete markers found (correct)")

print("\nSections That Should Be REMOVED:")
removed_correctly = []
not_removed = []
for section in should_be_removed:
    if section in sections_found:
        not_removed.append(section)
        idx, title = sections_found[section]
        print(f"  ❌ {section} still exists (line {idx}): {title}")
    else:
        removed_correctly.append(section)
        print(f"  ✓ {section} removed correctly")

print("\nSections That Should Be KEPT:")
kept_correctly = []
incorrectly_removed = []
for section in should_be_kept:
    if section in sections_found:
        kept_correctly.append(section)
        idx, title = sections_found[section]
        print(f"  ✓ {section} exists (line {idx}): {title}")
    else:
        incorrectly_removed.append(section)
        print(f"  ❌ {section} was removed but should be kept!")

print("\n" + "="*70)
print("SUMMARY")
print("="*70)

total_issues = len(delete_markers_found) + len(not_removed) + len(incorrectly_removed)

if total_issues == 0:
    print("✅ All sections handled correctly!")
    print(f"   - {len(removed_correctly)} sections removed")
    print(f"   - {len(kept_correctly)} sections kept")
else:
    print(f"⚠️  Found {total_issues} issue(s):")
    if delete_markers_found:
        print(f"   - {len(delete_markers_found)} delete markers still present")
    if not_removed:
        print(f"   - {len(not_removed)} sections not removed: {', '.join(not_removed)}")
    if incorrectly_removed:
        print(f"   - {len(incorrectly_removed)} sections incorrectly removed: {', '.join(incorrectly_removed)}")

# Also check for placeholder values
print("\n" + "="*70)
print("CHECKING PLACEHOLDER REPLACEMENT")
print("="*70)

placeholders_found = []
for para in doc.paragraphs:
    text = para.text
    if '{' in text and '}' in text:
        # Find all placeholders
        matches = re.findall(r'\{[^}]+\}', text)
        for match in matches:
            if match not in placeholders_found:
                placeholders_found.append(match)

if placeholders_found:
    print("⚠️  Unreplaced placeholders found:")
    for ph in placeholders_found:
        print(f"   - {ph}")
else:
    print("✓ All placeholders replaced")

print("\n" + "="*70)
