"""
Enhanced Letter of Claim Generator - FIXED VERSION
Handles placeholders split across multiple Word runs
"""

import json
import os
import sys
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Inches
from typing import List, Dict, Any, Optional
import argparse
import re


# Try to import config, fall back to defaults if not available
try:
    from .config import (
        DEFAULT_OUTPUT_DIR, 
        DEFENDANT_ADDRESSES,
        AGREEMENT_TYPES
    )
except ImportError:
    DEFAULT_OUTPUT_DIR = 'claim_letters'
    DEFENDANT_ADDRESSES = {}
    AGREEMENT_TYPES = {}


class ClaimLetterGenerator:
    """Generator for Letters of Claim from credit report JSON data."""
    
    def __init__(self, template_path: str):
        """
        Initialize the generator.
        
        Args:
            template_path: Path to the Word template file
        """
        self.template_path = template_path
        self.stats = {
            'total_reports': 0,
            'total_letters': 0,
            'by_client': {}
        }
        
        # Define conditional section markers
        self.conditional_sections = {
            '22.1': 'Additional Loan & Credit Agreement Commitments',
            '22.2': 'Personal Borrowing',
            '23.1': 'Personal Insolvency/Debt Management',
            '23.2': 'County Court Judgments (CCJs)',
            '23.3': 'Defaults',
            '23.4': 'Arrears',
            '23.5': 'Gambling',
            '23.6': 'Overdraft Usage'
        }
    
    @staticmethod
    def parse_client_name(full_name: str) -> Dict[str, str]:
        """Parse full name into first name and surname."""
        parts = full_name.strip().split()
        if len(parts) >= 2:
            return {
                'first_name': parts[0],
                'surname': ' '.join(parts[1:])
            }
        return {
            'first_name': full_name,
            'surname': ''
        }
    
    @staticmethod
    def parse_address(address: str) -> Dict[str, str]:
        """
        Parse address into components.
        
        Handles various address formats:
        - "00164 GWYNEDD AVENUE, L, SA1 6LL\nSWANSEA"
        - "123 High Street\nLondon\nSW1A 1AA"
        """
        if not address:
            return {
                'line1': '',
                'line2': '',
                'line3': '',
                'postcode': ''
            }
        
        # Split by newlines first
        lines = [line.strip() for line in address.split('\n') if line.strip()]
        
        postcode = ''
        address_lines = []
        
        # Process each line to extract postcode
        for i, line in enumerate(lines):
            # Check if this line contains a UK postcode pattern
            # UK postcode is usually at the end and has specific format
            parts = line.split(',')
            
            # Check the last part for postcode
            for j, part in enumerate(parts):
                part = part.strip()
                # UK postcode pattern: letters and numbers, typically 6-8 chars
                if len(part) >= 5 and len(part) <= 8:
                    # Check if it matches postcode pattern (has both letters and numbers)
                    has_letters = any(c.isalpha() for c in part)
                    has_numbers = any(c.isdigit() for c in part)
                    has_space = ' ' in part
                    
                    if has_letters and has_numbers:
                        postcode = part
                        # Rebuild line without postcode
                        remaining_parts = parts[:j] + parts[j+1:]
                        if remaining_parts:
                            line = ', '.join([p.strip() for p in remaining_parts if p.strip()])
                        else:
                            line = ''
                        break
            
            if line:
                address_lines.append(line)
        
        # If we didn't extract postcode yet, check the last line
        if not postcode and address_lines:
            last_line = address_lines[-1]
            words = last_line.split()
            
            # Check if last word(s) look like postcode
            for i in range(len(words) - 1, max(len(words) - 3, -1), -1):
                potential_postcode = ' '.join(words[i:])
                if len(potential_postcode) >= 5 and len(potential_postcode) <= 8:
                    has_letters = any(c.isalpha() for c in potential_postcode)
                    has_numbers = any(c.isdigit() for c in potential_postcode)
                    if has_letters and has_numbers:
                        postcode = potential_postcode
                        address_lines[-1] = ' '.join(words[:i])
                        if not address_lines[-1]:
                            address_lines.pop()
                        break
        
        # Ensure we have exactly 3 address lines
        while len(address_lines) < 3:
            address_lines.append('')
        
        # If we have more than 3 lines, combine them intelligently
        if len(address_lines) > 3:
            # Keep first line, combine middle lines, keep last line
            address_lines = [
                address_lines[0],
                ', '.join(address_lines[1:-1]),
                address_lines[-1]
            ]
        
        return {
            'line1': address_lines[0] if len(address_lines) > 0 else '',
            'line2': address_lines[1] if len(address_lines) > 1 else '',
            'line3': address_lines[2] if len(address_lines) > 2 else '',
            'postcode': postcode
        }
    
    @staticmethod
    def extract_bank_details_from_json(credit_data: Dict[str, Any]) -> Dict[str, str]:
        """
        Extract bank details from the JSON credit report.
        
        Args:
            credit_data: Full credit report data
            
        Returns:
            Dictionary with bank details
        """
        bank_details = {
            'bank_name': '',
            'account_name': '',
            'account_number': '',
            'sort_code': ''
        }
        
        # Try to get from client_info
        client_info = credit_data.get('credit_analysis', {}).get('client_info', {})
        
        # Extract bank name from bank_details if present
        json_bank_details = client_info.get('bank_details', {})
        if isinstance(json_bank_details, dict):
            bank_details['bank_name'] = json_bank_details.get('bank_name', '')
            bank_details['account_number'] = json_bank_details.get('account_number', '')
            bank_details['sort_code'] = json_bank_details.get('sort_code', '')
        
        # Account name should be the client's name
        client_name = client_info.get('name', '')
        if client_name:
            bank_details['account_name'] = client_name
        
        return bank_details
    
    @staticmethod
    def extract_account_details_from_lender(in_scope_item: Dict[str, Any]) -> Dict[str, str]:
        """
        Extract account number and start date from in-scope lender data.
        """
        account_details = {
            'account_number': 'TBC',
            'start_date': 'TBC'
        }
        
        # Handle account number - check single value first
        if 'account_number' in in_scope_item:
            account_number = in_scope_item.get('account_number')
            if account_number and str(account_number).strip():
                account_details['account_number'] = str(account_number).strip()
        
        # If single value wasn't available, try the array
        if account_details['account_number'] == 'TBC' and 'account_numbers' in in_scope_item:
            account_numbers = in_scope_item.get('account_numbers', [])
            if account_numbers and len(account_numbers) > 0:
                first_account = str(account_numbers[0]).strip()
                if first_account:
                    account_details['account_number'] = first_account
        
        # Handle start date - check single value first
        if 'start_date' in in_scope_item:
            start_date = in_scope_item.get('start_date')
            if start_date and str(start_date).strip():
                account_details['start_date'] = str(start_date).strip()
        
        # If single value wasn't available, try the array
        if account_details['start_date'] == 'TBC' and 'start_dates' in in_scope_item:
            start_dates = in_scope_item.get('start_dates', [])
            if start_dates and len(start_dates) > 0:
                first_date = str(start_dates[0]).strip()
                if first_date:
                    account_details['start_date'] = first_date
        
        return account_details
    
    @staticmethod
    def get_defendant_address(defendant_name: str, in_scope_item: Dict[str, Any]) -> str:
        """
        Get defendant address from config or JSON with flexible matching.
        
        Args:
            defendant_name: Name of the defendant
            in_scope_item: The in-scope lender data
            
        Returns:
            Address string
        """
        # First try exact match (case-sensitive)
        if defendant_name in DEFENDANT_ADDRESSES:
            return DEFENDANT_ADDRESSES[defendant_name]
        
        # Try exact match (case-insensitive)
        defendant_upper = defendant_name.upper()
        for key, address in DEFENDANT_ADDRESSES.items():
            if key.upper() == defendant_upper:
                return address
        
        # Try partial match - check if any config key is contained in the defendant name
        # This handles cases like "MONEYBARN NO" matching "Moneybarn No. 1 Limited"
        for key, address in DEFENDANT_ADDRESSES.items():
            # Normalize both strings for comparison
            key_normalized = key.upper().replace('.', '').replace(',', '')
            name_normalized = defendant_name.upper().replace('.', '').replace(',', '')
            
            # Check if the key is a substring of the defendant name
            if key_normalized in name_normalized:
                return address
            
            # Also try the reverse - check if defendant name starts with the key
            # This handles "VANQUIS BANK" matching "Vanquis Bank Ltd"
            if name_normalized.startswith(key_normalized):
                return address
        
        # Try to get from JSON if available
        if 'address' in in_scope_item:
            return in_scope_item['address']
        
        # Default
        return 'Address TBC'
    
    @staticmethod
    def extract_claim_metrics(credit_data: Dict[str, Any], target_lender_name: str = None) -> Dict[str, Any]:
        """
        Extract metrics needed for claim letter from credit analysis.
        
        DATA SOURCES:
        - FROM CREDIT REPORT: CCJs, Defaults, Arrears, AP markers
        - FROM BANK STATEMENTS (future): Gambling, Overdraft, P2P lending, Income/Expenditure
        
        Args:
            credit_data: Full credit report data
            target_lender_name: Name of the lender being claimed against (optional)
            
        Returns:
            Dictionary with calculated metrics and flags for conditional sections.
            Sections without data will be marked as unavailable and removed from letter.
        """
        credit_analysis = credit_data.get('credit_analysis', {})
        indicators = credit_analysis.get('indicators', {})
        claims_analysis = credit_analysis.get('claims_analysis', {})
        
        # Get reference date (use current date or agreement date if available)
        reference_date = datetime.now()
        twelve_months_ago = reference_date - timedelta(days=365)
        
        metrics = {
            # CCJ data
            'has_ccjs': False,
            'total_ccjs': 0,
            'ccj_details': [],
            
            # Defaults data
            'has_defaults': False,
            'totalDefaults12Months': 0,
            'defaults_details': [],
            
            # Arrears data
            'has_arrears': False,
            'totalArrears12Months': 0,
            'arrears_details': [],
            
            # AP marker (Personal Insolvency/Debt Management)
            'has_ap_marker': False,
            
            # Account commitments (estimated from credit limits)
            'has_loan_commitments': False,
            'totalContribution': 'N/A',
            'averageTotalContribution': 'N/A',
            
            # Data not available from credit report alone
            'has_peer_to_peer': False,
            'totalPeerToPeer': 'N/A',
            'has_gambling': False,
            'numberofTotalTransactions': 'N/A',
            'averageTotalGambling': 'N/A',
            'has_overdraft': False,
            'averageOverdraftUsageInDays': 'N/A',
            
            # Hard search flag
            'no_hard_search': False
        }
        
        # Extract CCJs from timeline
        credit_timeline = claims_analysis.get('credit_timeline', {})
        ccjs = credit_timeline.get('ccjs', [])
        
        if ccjs and len(ccjs) > 0:
            metrics['has_ccjs'] = True
            metrics['total_ccjs'] = len(ccjs)
            metrics['ccj_details'] = ccjs
        
        # Extract defaults from timeline
        defaults = credit_timeline.get('defaults', [])
        if defaults:
            metrics['has_defaults'] = True
            # Count defaults in last 12 months if date available
            for default in defaults:
                date_str = default.get('date', '')
                if date_str:
                    try:
                        default_date = datetime.strptime(date_str, '%d/%m/%Y')
                        if default_date >= twelve_months_ago:
                            metrics['totalDefaults12Months'] += 1
                    except:
                        # If we can't parse the date, count it anyway
                        metrics['totalDefaults12Months'] += 1
                else:
                    metrics['totalDefaults12Months'] += 1
            
            metrics['defaults_details'] = defaults
            
            # If no defaults in last 12 months but we have defaults, count them all
            if metrics['totalDefaults12Months'] == 0:
                metrics['totalDefaults12Months'] = len(defaults)
        
        # Check for defaults indicator
        if indicators.get('active_default', {}).get('flagged', False):
            metrics['has_defaults'] = True
            if metrics['totalDefaults12Months'] == 0:
                # Estimate at least 1 if indicator is flagged
                metrics['totalDefaults12Months'] = 1
        
        # Extract arrears from timeline
        arrears_pattern = credit_timeline.get('arrears_pattern', [])
        if arrears_pattern:
            metrics['has_arrears'] = True
            # Count total arrears months
            total_arrears_months = sum(
                item.get('arrears_months', 0) 
                for item in arrears_pattern
            )
            metrics['totalArrears12Months'] = total_arrears_months
            metrics['arrears_details'] = arrears_pattern
        
        # Check arrears indicator
        if indicators.get('arrears_last_6_months', {}).get('flagged', False):
            metrics['has_arrears'] = True
            if metrics['totalArrears12Months'] == 0:
                # Estimate 3 months if flagged
                metrics['totalArrears12Months'] = 3
        
        # Check for AP marker (arrangement to pay)
        if indicators.get('ap_marker', {}).get('flagged', False):
            metrics['has_ap_marker'] = True
        
        # Try to estimate loan commitments from in-scope and out-of-scope accounts
        in_scope = claims_analysis.get('in_scope', [])
        out_of_scope = claims_analysis.get('out_of_scope', [])
        
        total_credit_limit = 0
        account_count = 0
        
        for account in in_scope + out_of_scope:
            # Try to extract credit limit or loan value from account body/details
            # This is a rough estimation - would need more detailed account data
            account_count += 1
            # For now, we'll skip this as we don't have detailed balance data
        
        # Note: For peer-to-peer, gambling, and overdraft - these require bank statement analysis
        # which is not available from credit report data alone
        
        return metrics
    
    @staticmethod
    def remove_conditional_sections(doc: Document, metrics: Dict[str, Any]) -> None:
        """
        Remove conditional sections from document based on available metrics.
        
        Searches for section markers like "22.1 Additional Loan" and removes
        the entire paragraph if the required data is not available.
        
        Args:
            doc: Word document object
            metrics: Dictionary with metrics and availability flags
        """
        sections_to_remove = set()
        
        # Determine which sections to remove (just the numbers)
        if not metrics.get('has_loan_commitments', False):
            sections_to_remove.add('22.1')
        
        if not metrics.get('has_peer_to_peer', False):
            sections_to_remove.add('22.2')
        
        if not metrics.get('has_ap_marker', False):
            sections_to_remove.add('23.1')
        
        if not metrics.get('has_ccjs', False):
            sections_to_remove.add('23.2')
        
        if not metrics.get('has_defaults', False):
            sections_to_remove.add('23.3')
        
        if not metrics.get('has_arrears', False):
            sections_to_remove.add('23.4')
        
        if not metrics.get('has_gambling', False):
            sections_to_remove.add('23.5')
        
        if not metrics.get('has_overdraft', False):
            sections_to_remove.add('23.6')
        
        # Track paragraphs to remove
        paragraphs_to_remove = []
        
        # Iterate through paragraphs to find section markers
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            para_text = para.text.strip()
            
            # Check for delete markers - always remove these
            if '{*Delete' in para_text and ('Not Applicable' in para_text or 'Not Appliable' in para_text):
                # This is a conditional marker - check if we should keep the sections below
                # For now, always remove these markers
                paragraphs_to_remove.append(para)
                i += 1
                continue
            
            # Check if this paragraph starts with a section marker we want to remove
            # Handle format like "22.1 " or "       22.1 " (with any leading whitespace)
            should_remove_section = False
            matched_section = None
            
            for section_num in sections_to_remove:
                # Try to match with various patterns:
                # "22.1 Text" or "       22.1 Text" (with leading spaces)
                pattern = r'^\s*' + re.escape(section_num) + r'\s+'
                if re.match(pattern, para_text):
                    should_remove_section = True
                    matched_section = section_num
                    break
            
            if should_remove_section:
                # Mark this paragraph for removal
                paragraphs_to_remove.append(para)
                i += 1
                
                # Continue removing paragraphs until we hit another numbered section
                # or significant break
                while i < len(doc.paragraphs):
                    next_para = doc.paragraphs[i]
                    next_text = next_para.text.strip()
                    
                    # Stop if we hit another section marker (XX.Y format)
                    # Like "22.1", "22.2", "23.1", etc.
                    if re.match(r'^\d{2}\.\d+\s', next_text):
                        break
                    
                    # Stop if we hit a major section marker (XX. format)
                    # Like "20.", "21.", "22.", "23." (but not subsections)
                    if re.match(r'^\d{2}\.\s+[A-Z]', next_text):
                        break
                    
                    # Stop at empty paragraphs that might indicate section end
                    if not next_text:
                        # Look ahead - if next non-empty is a section, stop here
                        look_ahead = i + 1
                        while look_ahead < len(doc.paragraphs):
                            future_text = doc.paragraphs[look_ahead].text.strip()
                            if future_text:
                                if re.match(r'^\d{2}\.\d*\s', future_text):
                                    # Next real para is a section - stop removal
                                    break
                                else:
                                    # Next para is content, include this empty para
                                    paragraphs_to_remove.append(next_para)
                                    i += 1
                                    break
                            look_ahead += 1
                        else:
                            # No more content, include empty para
                            paragraphs_to_remove.append(next_para)
                            i += 1
                        break
                    
                    # Otherwise, include this paragraph in removal
                    paragraphs_to_remove.append(next_para)
                    i += 1
            else:
                i += 1
        
        # Remove marked paragraphs in reverse order (to avoid index issues)
        for para in reversed(paragraphs_to_remove):
            p_element = para._element
            p_element.getparent().remove(p_element)
    
    @staticmethod
    def replace_text_in_paragraph(paragraph, search_str: str, replace_str: str) -> bool:
        """
        Replace text in a paragraph, handling text split across multiple runs.
        This is the key fix - it reassembles the full text, does the replacement,
        then puts it back in the first run.
        
        Args:
            paragraph: The paragraph object
            search_str: Text to search for (e.g., '{Account Number}')
            replace_str: Text to replace with
            
        Returns:
            bool: True if replacement was made
        """
        # Get the complete text from all runs
        full_text = ''.join(run.text for run in paragraph.runs)
        
        # Check if search string exists (case-sensitive, but handles leading/trailing whitespace)
        if search_str in full_text or search_str.strip() in full_text:
            # Do the replacement
            new_text = full_text.replace(search_str, replace_str)
            
            # If that didn't work, try trimming the search string
            if new_text == full_text:
                new_text = full_text.replace(search_str.strip(), replace_str)
            
            # Put all the text in the first run, clear the others
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ''
                return True
        
        return False
    
    @staticmethod
    def replace_text_in_cell(cell, search_str: str, replace_str: str) -> bool:
        """Replace text in a table cell - handles multiple paragraphs and nested tables."""
        replaced = False
        
        # Replace in each paragraph
        for paragraph in cell.paragraphs:
            if ClaimLetterGenerator.replace_text_in_paragraph(paragraph, search_str, replace_str):
                replaced = True
        
        # Also check nested tables
        for table in cell.tables:
            if ClaimLetterGenerator.replace_text_in_table(table, search_str, replace_str):
                replaced = True
        
        return replaced
    
    @staticmethod
    def replace_text_in_table(table, search_str: str, replace_str: str) -> bool:
        """Replace text in all cells of a table."""
        replaced = False
        for row in table.rows:
            for cell in row.cells:
                if ClaimLetterGenerator.replace_text_in_cell(cell, search_str, replace_str):
                    replaced = True
        return replaced
    
    def _inject_logo(self, doc: Document, logo_path: str, width_inches: float = 2.0):
        """Replace the {Company Logo} placeholder with the company logo image.

        Searches body paragraphs then header/footer areas.  The first paragraph
        containing the placeholder is cleared and an inline picture is inserted.
        """
        target = '{Company Logo}'

        def _try_paragraphs(paragraphs):
            for para in paragraphs:
                if target in para.text:
                    para.clear()
                    para.add_run().add_picture(logo_path, width=Inches(width_inches))
                    return True
            return False

        if _try_paragraphs(doc.paragraphs):
            return
        for section in doc.sections:
            for zone in [section.header, section.first_page_header,
                         section.footer, section.first_page_footer]:
                if zone and _try_paragraphs(zone.paragraphs):
                    return

    def replace_placeholders(self, doc: Document, replacements: Dict[str, str]):
        """
        Replace all placeholders in the document.
        Goes through each placeholder and replaces it everywhere in the document.
        """
        for search_text, replace_text in replacements.items():
            replace_str = str(replace_text) if replace_text is not None else ''
            
            # Replace in main document paragraphs
            for paragraph in doc.paragraphs:
                self.replace_text_in_paragraph(paragraph, search_text, replace_str)
            
            # Replace in all tables
            for table in doc.tables:
                self.replace_text_in_table(table, search_text, replace_str)
            
            # Replace in headers
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for paragraph in header.paragraphs:
                            self.replace_text_in_paragraph(paragraph, search_text, replace_str)
                        for table in header.tables:
                            self.replace_text_in_table(table, search_text, replace_str)
                
                # Replace in footers
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for paragraph in footer.paragraphs:
                            self.replace_text_in_paragraph(paragraph, search_text, replace_str)
                        for table in footer.tables:
                            self.replace_text_in_table(table, search_text, replace_str)
    
    def generate_letter(self, output_path: str, credit_data: Dict[str, Any],
                       in_scope_item: Dict[str, Any], debug: bool = False,
                       branding: Optional[Dict[str, Any]] = None) -> bool:
        """
        Generate a single letter of claim.
        
        Args:
            output_path: Where to save the generated letter
            credit_data: Full credit report data
            in_scope_item: The specific in-scope lender data
            debug: If True, print debug information
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Load template
            doc = Document(self.template_path)
            
            # Parse client information
            client_info = credit_data['credit_analysis']['client_info']
            client_name_parts = self.parse_client_name(client_info['name'])
            client_address = self.parse_address(client_info['address'])
            
            # Extract bank details from JSON
            bank_details = self.extract_bank_details_from_json(credit_data)
            
            # Extract account details from the in_scope_item
            account_details = self.extract_account_details_from_lender(in_scope_item)
            
            # Get defendant information
            defendant_name = in_scope_item['name']
            defendant_address = self.get_defendant_address(defendant_name, in_scope_item)
            
            # Extract claim metrics from credit analysis
            metrics = self.extract_claim_metrics(credit_data, defendant_name)
            
            # Debug: Print what we extracted
            if debug:
                print(f"    → Account Number: {account_details['account_number']}")
                print(f"    → Start Date: {account_details['start_date']}")
                print(f"    → Metrics extracted:")
                print(f"      - Has CCJs: {metrics['has_ccjs']} ({metrics['total_ccjs']})")
                print(f"      - Has Defaults: {metrics['has_defaults']} ({metrics['totalDefaults12Months']})")
                print(f"      - Has Arrears: {metrics['has_arrears']} ({metrics['totalArrears12Months']})")
                print(f"      - Has AP Marker: {metrics['has_ap_marker']}")
            
            # Remove conditional sections that don't have data
            self.remove_conditional_sections(doc, metrics)
            
            # Get current date
            current_date = datetime.now().strftime('%d/%m/%Y')
            
            # Prepare replacements dictionary with calculated metrics
            replacements = {
                '{Date}': current_date,
                '{Defendant Name}': defendant_name,
                '{Defendant Address}': defendant_address,
                '{Client First Name}': client_name_parts['first_name'],
                '{Client Surname}': client_name_parts['surname'],
                '{Address Line 1}': client_address['line1'],
                '{Address Line 2}': client_address['line2'],
                '{Address Line 3}': client_address['line3'],
                '{Postcode}': client_address['postcode'],
                '{Bank}': bank_details.get('bank_name', defendant_name),
                '{Account Name}': bank_details.get('account_name', client_info['name']),
                '{Account Number}': account_details['account_number'],
                '{Sort Code}': bank_details.get('sort_code', 'TBC'),
                '{Agreement Number}': account_details['account_number'],
                '{Agreement Start Date}': account_details['start_date'],
                '{Report Received Date}': current_date,
                '{Report Outcome}': 'unaffordable',
                # Income/Expenditure - not available from credit report
                '{averageConsistentIncome}': 'TBC',
                '{averageCommittedExpenditure}': 'TBC',
                '{averageLivingExpenditure}': 'TBC',
                '{totalAverageExpenditure}': 'TBC',
                '{disposableincome}': 'TBC',
                # Loan commitments - calculated from credit data
                '{totalContribution}': metrics.get('totalContribution', 'N/A'),
                '{averageTotalContribution}': metrics.get('averageTotalContribution', 'N/A'),
                # Peer-to-peer lending - requires bank statement analysis
                '{totalPeerToPeer}': metrics.get('totalPeerToPeer', 'N/A'),
                # Credit report metrics - calculated from analysis
                '{totalDefaults12Months}': str(metrics.get('totalDefaults12Months', 0)),
                '{totalArrears12Months}': str(metrics.get('totalArrears12Months', 0)),
                # Gambling - requires bank statement analysis
                '{numberofTotalTransactions}': metrics.get('numberofTotalTransactions', 'N/A'),
                '{averageTotalGambling}': metrics.get('averageTotalGambling', 'N/A'),
                # Overdraft - requires bank statement analysis
                '{averageOverdraftUsageInDays}': metrics.get('averageOverdraftUsageInDays', 'N/A'),
            }
            
            # Replace all placeholders
            self.replace_placeholders(doc, replacements)

            # Apply company branding
            if branding:
                footer_msg = branding.get('footer_message', '')
                if footer_msg:
                    self.replace_placeholders(doc, {'{Footer Message}': footer_msg})
                logo_path = branding.get('logo_path', '')
                if logo_path and os.path.exists(logo_path):
                    self._inject_logo(doc, logo_path)

            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"✗ Error generating {output_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def process_credit_report(self, report_data: Dict[str, Any], output_dir: str, debug: bool = False) -> int:
        """
        Process a single credit report and generate letters for all in-scope lenders.
        
        Args:
            report_data: Credit report JSON data
            output_dir: Directory to save generated letters
            debug: If True, print debug information
            
        Returns:
            int: Number of letters generated
        """
        credit_analysis = report_data.get('credit_analysis', {})
        client_info = credit_analysis.get('client_info', {})
        client_name = client_info.get('name', 'Unknown_Client')
        
        # Clean client name for filename
        safe_client_name = ''.join(c if c.isalnum() or c in [' ', '_'] else '_' for c in client_name)
        safe_client_name = safe_client_name.replace(' ', '_')
        
        claims_analysis = credit_analysis.get('claims_analysis', {})
        in_scope_lenders = claims_analysis.get('in_scope', [])
        
        letters_generated = 0
        
        if in_scope_lenders:
            print(f"\n📄 {client_name}: {len(in_scope_lenders)} in-scope lender(s)")
            
            for i, lender in enumerate(in_scope_lenders):
                lender_name = lender.get('name', 'Unknown_Lender')
                safe_lender_name = ''.join(c if c.isalnum() or c in [' ', '_'] else '_' for c in lender_name)
                safe_lender_name = safe_lender_name.replace(' ', '_')
                
                output_filename = f"{safe_client_name}_{safe_lender_name}_LOC.docx"
                output_path = os.path.join(output_dir, output_filename)
                
                # Only debug the first letter
                debug_this = debug and i == 0
                
                if self.generate_letter(output_path, report_data, lender, debug=debug_this):
                    print(f"  ✓ {lender_name}")
                    letters_generated += 1
                else:
                    print(f"  ✗ {lender_name} - Failed")
        else:
            print(f"\n⚠ {client_name}: No in-scope lenders found")
        
        # Update stats
        self.stats['by_client'][client_name] = letters_generated
        
        return letters_generated
    
    def load_json_data(self, json_input: str) -> List[Dict[str, Any]]:
        """
        Load JSON data from file, directory, or string.
        
        Handles the format: [{"url": "...", "credit_analysis": {...}}]
        
        Args:
            json_input: Path to file/directory or JSON string
            
        Returns:
            List of credit report dictionaries
        """
        json_files = []
        json_path = Path(json_input)
        
        if json_path.is_file() and json_path.suffix == '.json':
            with open(json_path, 'r') as f:
                data = json.load(f)
                json_files = data if isinstance(data, list) else [data]
        elif json_path.is_dir():
            for file_path in sorted(json_path.glob('*.json')):
                with open(file_path, 'r') as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        json_files.extend(data)
                    else:
                        json_files.append(data)
        else:
            try:
                data = json.loads(json_input)
                json_files = data if isinstance(data, list) else [data]
            except:
                raise ValueError(f"Invalid JSON input: {json_input}")
        
        return json_files
    
    def generate_all(self, json_input: str, output_dir: str, debug: bool = False) -> Dict[str, Any]:
        """
        Generate letters for all credit reports in the input.
        
        Args:
            json_input: Path to JSON file/directory or JSON string
            output_dir: Directory to save generated letters
            debug: If True, print debug information for first letter
            
        Returns:
            Dictionary with generation statistics
        """
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Load data
        print("📂 Loading credit report data...")
        credit_reports = self.load_json_data(json_input)
        self.stats['total_reports'] = len(credit_reports)
        print(f"   Found {len(credit_reports)} credit report(s)\n")
        
        # Process each report
        for i, report_data in enumerate(credit_reports):
            # Only debug first report
            debug_this = debug and i == 0
            letters_count = self.process_credit_report(report_data, output_dir, debug=debug_this)
            self.stats['total_letters'] += letters_count
        
        # Print summary
        self.print_summary(output_dir)
        
        return self.stats
    
    def print_summary(self, output_dir: str):
        """Print generation summary."""
        print(f"\n{'='*70}")
        print(f"📊 GENERATION SUMMARY")
        print(f"{'='*70}")
        print(f"Total Reports Processed: {self.stats['total_reports']}")
        print(f"Total Letters Generated: {self.stats['total_letters']}")
        print(f"Output Directory: {output_dir}")
        print(f"{'='*70}\n")


def main():
    """Main entry point for command-line usage."""
    parser = argparse.ArgumentParser(
        description='Generate Letters of Claim for in-scope lenders from credit report JSON',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process a single JSON file
  python generate_claim_letters_final.py data.json
  
  # Process with custom template and output directory
  python generate_claim_letters_final.py data.json -t template.docx -o output_letters
  
  # Enable debug mode
  python generate_claim_letters_final.py data.json --debug
        """
    )
    
    parser.add_argument('json_input', help='Path to JSON file, directory, or JSON string')
    parser.add_argument('-o', '--output', default=DEFAULT_OUTPUT_DIR,
                       help=f'Output directory (default: {DEFAULT_OUTPUT_DIR})')
    parser.add_argument('-t', '--template', 
                       default='Affordability_Letter_of_Claim_.docx',
                       help='Path to Word template')
    parser.add_argument('--debug', action='store_true',
                       help='Enable debug mode to see document structure')
    
    args = parser.parse_args()
    
    # Generate letters
    generator = ClaimLetterGenerator(args.template)
    generator.generate_all(args.json_input, args.output, debug=args.debug)


if __name__ == '__main__':
    main()